"""
DocQuery — Chat Export (Phase 3)

Export conversations as Markdown or PDF.
"""

import io
import re
from datetime import datetime
from typing import Literal, Optional, List, Dict, Any

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from src.api.dependencies import get_current_user, require_cap
from src.api.routes.audit import log_audit
from src.logger import get_logger

router = APIRouter()
logger = get_logger(__name__)


# ── Phase 4.3: XLSX export of Analyst tables (preserves computed columns) ──────

class TableExportRequest(BaseModel):
    """A table answer to export: an optional source grid + computed Analyst rows.

    `grid` is the normalized table (headers + rows) the answer drew from; `computed`
    is the list of Analyst results (operation/value/formula/sources). Both are kept
    in the workbook so the export preserves the computed columns and their formulas.
    """
    title: Optional[str] = "DocQuery Table Export"
    grid: Optional[Dict[str, Any]] = None        # {headers:[...], rows:[{...}]}
    computed: Optional[List[Dict[str, Any]]] = None  # analyst.results_to_rows() output


def _build_xlsx(req: "TableExportRequest") -> bytes:
    """Build an .xlsx with two sheets: the source grid and the computed results.

    Computed columns are clearly labelled and carry the shown formula + source-cell
    trace, so a number in the workbook still traces to a source cell or a formula
    (the §4b contract survives the export, not just the on-screen answer).
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    header_font = Font(bold=True)
    computed_fill = PatternFill("solid", fgColor="FFF3E0")  # tint computed rows

    # Sheet 1: the source table grid (raw cells, the authoritative numbers)
    ws = wb.active
    ws.title = "Source Table"
    grid = req.grid or {}
    headers = grid.get("headers", [])
    if headers:
        ws.append(headers)
        for c in ws[1]:
            c.font = header_font
        for row in grid.get("rows", []):
            ws.append([row.get(h, "") for h in headers])
    else:
        ws.append(["(no source grid provided)"])

    # Sheet 2: the computed results (value + formula + provenance)
    ws2 = wb.create_sheet("Computed")
    cols = ["operation", "value", "unit", "formula", "sources", "error"]
    ws2.append([c.title() for c in cols])
    for c in ws2[1]:
        c.font = header_font
    for r in (req.computed or []):
        ws2.append([r.get(c, "") for c in cols])
        for cell in ws2[ws2.max_row]:
            cell.fill = computed_fill

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── G6.1: .docx export of a cited deliverable (draft/memo/summary) ─────────────
#
# The agent emits cited markdown that ALREADY passed the output gate (every factual
# sentence carries a `[... p.N]` marker or was withheld). Export is dumb plumbing:
# markdown → .docx, PRESERVING the citation contract. It never re-gates and never
# adds a number. Two modes — "with citations" turns the inline markers into numbered
# endnotes (the lawyer's working copy); "without" strips them to a clean client copy.
# Critical: the gate ran upstream on the cited form; stripping happens here only.

# Same citation-marker shape the gate enforces (gates.py `_CITE_RE`): a bracketed
# token containing `p.N`, e.g. [doc p.3], [amzn-2022 p.41], [MSFT FY22 p.7].
_DOCX_CITE_RE = re.compile(r"\[[^\]]*\bp\.?\s*\d+[^\]]*\]")

# Inline markdown we flatten when writing a docx run: **bold**, *italic*, `code`.
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_MD_CODE_RE = re.compile(r"`([^`]+)`")


class DocxExportRequest(BaseModel):
    """A cited deliverable to export as .docx.

    `markdown` is the gated draft (headings + body + inline `[... p.N]` markers).
    `include_citations=True` renders the markers as numbered endnotes; `False`
    strips them for a clean client copy. The body numbers are IDENTICAL either way
    — only the provenance markers differ — because the gate already bound them.
    """
    title: Optional[str] = "DocQuery Draft"
    markdown: str = ""
    include_citations: bool = True
    # F1e: stamp the privilege watermark when this export contains privileged material. The
    # caller (the surface that owns the doc set) sets it; default False ⇒ no watermark (the
    # byte-identical pre-F1e export).
    privileged: bool = False


def _strip_inline_md(text: str):
    """Yield (text, fmt) runs from a line, flattening **bold**/*italic*/`code`.

    Returns a list of (segment, {'bold','italic','code'}) tuples so python-docx can
    apply run-level formatting. Citation markers are handled BEFORE this (they're
    pulled out as endnote refs upstream), so they don't appear here.
    """
    # Tokenize on the three inline markers, tagging each captured group.
    runs: List = []
    pos = 0
    # Combined scan: find the earliest of bold/code/italic at each step.
    pattern = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], {}))
        if m.group(1) is not None:
            runs.append((m.group(1), {"bold": True}))
        elif m.group(2) is not None:
            runs.append((m.group(2), {"code": True}))
        elif m.group(3) is not None:
            runs.append((m.group(3), {"italic": True}))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], {}))
    return runs or [(text, {})]


def _build_docx(req: "DocxExportRequest") -> bytes:
    """Render gated markdown to a .docx, preserving the citation contract.

    Markdown handled: ATX headings (`#`..`######`), unordered (`- `/`* `) and ordered
    (`1. `) list items, blank-line paragraph breaks, and inline **bold**/*italic*/`code`.
    Citation markers `[... p.N]` become superscript endnote refs (include_citations)
    collected into a "Sources" section, or are stripped (clean copy). No number is
    ever invented or removed — only the provenance markers move.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(req.title or "DocQuery Draft", level=0)

    # F1e: privilege watermark — a bold notice line at the top when the export contains
    # privileged material, so a privileged doc never leaves the system unmarked.
    if getattr(req, "privileged", False):
        from src.components.privilege import WATERMARK_NOTICE
        wm = doc.add_paragraph()
        wm.add_run(WATERMARK_NOTICE).bold = True

    endnotes: List[str] = []           # ordered unique citation markers
    endnote_index: Dict[str, int] = {}  # marker text -> 1-based endnote number

    def cite_number(marker: str) -> int:
        if marker not in endnote_index:
            endnotes.append(marker)
            endnote_index[marker] = len(endnotes)
        return endnote_index[marker]

    def add_runs(paragraph, text: str):
        """Write `text` into `paragraph`, lifting citation markers to endnote refs."""
        last = 0
        for cm in _DOCX_CITE_RE.finditer(text):
            # plain text before the marker (with inline md), then the marker itself
            for seg, fmt in _strip_inline_md(text[last:cm.start()]):
                r = paragraph.add_run(seg)
                if fmt.get("bold"): r.bold = True
                if fmt.get("italic"): r.italic = True
                if fmt.get("code"): r.font.name = "Courier New"
            if req.include_citations:
                marker = cm.group(0)
                ref = paragraph.add_run(f"[{cite_number(marker)}]")
                ref.font.superscript = True
            # without citations: marker is simply dropped
            last = cm.end()
        for seg, fmt in _strip_inline_md(text[last:]):
            r = paragraph.add_run(seg)
            if fmt.get("bold"): r.bold = True
            if fmt.get("italic"): r.italic = True
            if fmt.get("code"): r.font.name = "Courier New"

    para_buf: List[str] = []

    def flush_paragraph():
        if not para_buf:
            return
        text = " ".join(para_buf).strip()
        para_buf.clear()
        if text:
            add_runs(doc.add_paragraph(), text)

    for raw in (req.markdown or "").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            flush_paragraph()
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            flush_paragraph()
            add_runs(doc.add_heading(level=min(len(h.group(1)), 6)), h.group(2).strip())
            continue
        ul = re.match(r"^\s*[-*]\s+(.*)$", line)
        if ul:
            flush_paragraph()
            add_runs(doc.add_paragraph(style="List Bullet"), ul.group(1).strip())
            continue
        ol = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if ol:
            flush_paragraph()
            add_runs(doc.add_paragraph(style="List Number"), ol.group(1).strip())
            continue
        para_buf.append(line.strip())
    flush_paragraph()

    # Endnotes / sources section (only when citations are kept)
    if req.include_citations and endnotes:
        doc.add_heading("Sources", level=1)
        for i, marker in enumerate(endnotes, start=1):
            p = doc.add_paragraph()
            p.add_run(f"[{i}] ").bold = True
            p.add_run(marker.strip("[]"))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── G6.3: tracked-changes .docx for a redline result ────────────────────────────
#
# Each RedlineFinding with status="deviation" becomes a section in the .docx:
#   - The target clause quote (what the contract says) as a struck-through "deleted" run.
#   - The suggested edit as an underlined "inserted" run.
#   - The rationale as a comment-style paragraph.
# Conforming and missing clauses are noted without tracked changes.
# python-docx supports run-level formatting (bold, underline, strike) — we use those
# to represent tracked changes visually (Word opens them as a clean redlined document).
# True OOXML revision marks (w:ins / w:del) would require low-level XML; the
# formatting approach is lawyer-readable and sufficient for the India mid-market wedge.

class RedlineExportRequest(BaseModel):
    """A completed redline result to export as a tracked-changes .docx."""
    title: Optional[str] = "Redline Review"
    doc_name: Optional[str] = None
    findings: List[Dict[str, Any]] = []   # list of RedlineFinding-shaped dicts
    # F1e: stamp the privilege watermark when the redlined document is privileged.
    privileged: bool = False


def _build_redline_docx(req: "RedlineExportRequest") -> bytes:
    """Render a redline result to a tracked-changes .docx.

    Each deviation finding gets a section with: the original clause (struck-through)
    and the suggested replacement (underlined + bold). Conforming/missing/abstain
    clauses are recorded as summary lines. No figure is invented; export is dumb
    plumbing — the agent already grounded the findings.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document()
    doc.add_heading(req.title or "Redline Review", level=0)
    # F1e: privilege watermark at the top when the redlined doc is privileged.
    if getattr(req, "privileged", False):
        from src.components.privilege import WATERMARK_NOTICE
        wm = doc.add_paragraph()
        wm.add_run(WATERMARK_NOTICE).bold = True
    if req.doc_name:
        p = doc.add_paragraph()
        p.add_run(f"Document: {req.doc_name}").italic = True
    doc.add_paragraph()

    deviations = [f for f in req.findings if f.get("status") == "deviation"]
    conforming = [f for f in req.findings if f.get("status") == "conforming"]
    missing = [f for f in req.findings if f.get("status") == "missing"]
    abstained = [f for f in req.findings if f.get("status") == "abstain"]

    # ── Deviations (the main redline content) ──
    if deviations:
        doc.add_heading("Deviations", level=1)
        for f in deviations:
            doc.add_heading(f.get("clause_topic", "Clause"), level=2)

            # Original clause — struck-through (the "deleted" tracked change)
            if f.get("target_quote"):
                p = doc.add_paragraph()
                p.add_run("Current: ").bold = True
                r = p.add_run(f["target_quote"])
                r.font.strike = True
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            # Suggested edit — underlined + green (the "inserted" tracked change)
            if f.get("suggested_edit"):
                p = doc.add_paragraph()
                p.add_run("Suggested: ").bold = True
                r = p.add_run(f["suggested_edit"])
                r.underline = True
                r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)

            # Deviation description
            if f.get("deviation"):
                p = doc.add_paragraph()
                p.add_run("Issue: ").bold = True
                p.add_run(f["deviation"])

            # Rationale (cites the playbook rule)
            if f.get("rationale"):
                p = doc.add_paragraph()
                p.add_run("Rationale: ").bold = True
                r = p.add_run(f["rationale"])
                r.italic = True

            # Firm standard position for reference
            if f.get("playbook_standard"):
                p = doc.add_paragraph()
                p.add_run("Firm standard: ").bold = True
                r = p.add_run(f["playbook_standard"])
                r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                r.italic = True

            doc.add_paragraph()

    # ── Conforming clauses (summary only) ──
    if conforming:
        doc.add_heading("Conforming Clauses", level=1)
        for f in conforming:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f.get("clause_topic", "Clause")).bold = True
            if f.get("target_quote"):
                p.add_run(f" — {f['target_quote'][:120]}{'…' if len(f.get('target_quote','')) > 120 else ''}")

    # ── Missing clauses ──
    if missing:
        doc.add_heading("Missing Clauses", level=1)
        for f in missing:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f.get("clause_topic", "Clause")).bold = True
            p.add_run(" — clause not found in document")
            if f.get("playbook_standard"):
                doc.add_paragraph(f"  Firm standard: {f['playbook_standard'][:200]}")

    # ── Abstained (could not assess) ──
    if abstained:
        doc.add_heading("Could Not Assess", level=1)
        for f in abstained:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f.get("clause_topic", "Clause")).bold = True
            if f.get("rationale"):
                p.add_run(f" — {f['rationale']}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── G6.1b: .pdf export of a cited deliverable (fast-follow, fpdf2) ──────────────
#
# Same contract as .docx: markdown ALREADY passed the gate; export is dumb plumbing.
# fpdf2 is already pinned (used by the conversation-PDF endpoint). weasyprint requires
# native GTK/Pango libs not present on all platforms — fpdf2 is pure-Python, zero
# new system deps. Citation markers → numbered footnotes (with) or stripped (without).

class PdfExportRequest(BaseModel):
    """A gated cited deliverable to export as .pdf."""
    title: Optional[str] = "DocQuery Draft"
    markdown: str = ""
    include_citations: bool = True
    # F1e: stamp the privilege watermark when this export contains privileged material.
    privileged: bool = False


def _build_pdf(req: "PdfExportRequest") -> bytes:
    """Render gated markdown to a .pdf using fpdf2, preserving the citation contract."""
    from fpdf import FPDF

    class _PDF(FPDF):
        def header(self):
            pass  # suppress default header

    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    endnotes: List[str] = []
    endnote_index: Dict[str, int] = {}

    def cite_number(marker: str) -> int:
        if marker not in endnote_index:
            endnotes.append(marker)
            endnote_index[marker] = len(endnotes)
        return endnote_index[marker]

    def _safe(text: str) -> str:
        return text.encode("latin-1", "replace").decode("latin-1")

    def render_line(text: str, *, bold: bool = False, size: int = 10, indent: int = 0):
        """Write one text line, lifting citation markers to inline refs."""
        # Collect segments: plain text or citation markers
        parts: List = []
        last = 0
        for cm in _DOCX_CITE_RE.finditer(text):
            if cm.start() > last:
                parts.append(("text", text[last:cm.start()]))
            if req.include_citations:
                parts.append(("cite", f"[{cite_number(cm.group(0))}]"))
            last = cm.end()
        if last < len(text):
            parts.append(("text", text[last:]))

        # Strip markdown inline markers from plain segments
        _inline = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")

        x0 = pdf.l_margin + indent
        pdf.set_x(x0)
        for kind, segment in parts:
            if kind == "cite":
                pdf.set_font("Helvetica", "I", size - 2)
                pdf.write(size * 0.4, _safe(segment))
                pdf.set_font("Helvetica", "B" if bold else "", size)
            else:
                # Flatten inline md
                clean = _inline.sub(lambda m: (m.group(1) or m.group(2) or m.group(3) or ""), segment)
                pdf.set_font("Helvetica", "B" if bold else "", size)
                pdf.write(size * 0.4, _safe(clean))

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, _safe(req.title or "DocQuery Draft"), ln=True)
    pdf.ln(2)

    # F1e: privilege watermark — a bold notice under the title when privileged.
    if getattr(req, "privileged", False):
        from src.components.privilege import WATERMARK_NOTICE
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(176, 38, 38)  # red — a trust/withheld signal, design-system aligned
        pdf.multi_cell(0, 5, _safe(WATERMARK_NOTICE))
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    for raw in (req.markdown or "").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            pdf.ln(4)
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            depth = len(h.group(1))
            size = max(10, 15 - depth)
            pdf.set_font("Helvetica", "B", size)
            pdf.ln(3)
            render_line(h.group(2).strip(), bold=True, size=size)
            pdf.ln(size * 0.5)
            continue
        ul = re.match(r"^\s*[-*]\s+(.*)$", line)
        if ul:
            render_line("• " + ul.group(1).strip(), indent=4)
            pdf.ln(5)
            continue
        ol = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if ol:
            render_line(f"{ol.group(1)}. {ol.group(2).strip()}", indent=4)
            pdf.ln(5)
            continue
        render_line(line.strip())
        pdf.ln(5)

    # Sources section
    if req.include_citations and endnotes:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Sources", ln=True)
        pdf.set_font("Helvetica", "", 9)
        for i, marker in enumerate(endnotes, start=1):
            pdf.multi_cell(0, 5, _safe(f"[{i}] {marker.strip('[]')}"))

    return bytes(pdf.output())


def _format_conversation_md(title: str, messages: list) -> str:
    """Format a conversation as Markdown text."""
    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"*Exported on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}*")
    lines.append("")
    lines.append("---")
    lines.append("")

    for msg in messages:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        created = msg.get("created_at", "")
        sources = msg.get("sources", [])

        if role == "user":
            lines.append(f"## 🧑 You")
        else:
            lines.append(f"## 🤖 DocQuery")

        if created:
            try:
                dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
                lines.append(f"*{dt.strftime('%I:%M %p')}*")
            except Exception:
                pass

        lines.append("")
        lines.append(content)
        lines.append("")

        # Add sources for assistant messages
        if sources and role == "assistant":
            lines.append("**Sources:**")
            for s in sources:
                filename = s.get("filename", "Unknown")
                page = s.get("page", "")
                page_str = f" (p. {page})" if page else ""
                lines.append(f"- {filename}{page_str}")
            lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def _format_conversation_pdf(title: str, messages: list) -> bytes:
    """Format a conversation as a PDF document using fpdf2."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF export requires fpdf2. Install with: pip install fpdf2"
        )

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, title, ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 8, f"Exported on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    for msg in messages:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        sources = msg.get("sources", [])

        # Role header
        pdf.set_font("Helvetica", "B", 12)
        label = "You" if role == "user" else "DocQuery"
        pdf.set_fill_color(240, 240, 240) if role == "user" else pdf.set_fill_color(245, 245, 255)
        pdf.cell(0, 8, label, ln=True, fill=True)

        # Content
        pdf.set_font("Helvetica", "", 10)
        pdf.ln(2)

        # Handle multi-line content safely
        for line in content.split("\n"):
            # fpdf2 multi_cell handles wrapping
            pdf.multi_cell(0, 5, line.encode("latin-1", "replace").decode("latin-1"))

        # Sources
        if sources and role == "assistant":
            pdf.ln(2)
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(100, 100, 100)
            for s in sources:
                filename = s.get("filename", "Unknown")
                page = s.get("page", "")
                page_str = f" (p. {page})" if page else ""
                pdf.cell(0, 5, f"Source: {filename}{page_str}", ln=True)
            pdf.set_text_color(0, 0, 0)

        pdf.ln(4)

        # Separator line
        pdf.set_draw_color(200, 200, 200)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(4)

    return pdf.output()


@router.post("/export/pdf")
async def export_draft_pdf(
    req: PdfExportRequest,
    sb=Depends(get_current_user),
    _cap=Depends(require_cap("release_external")),
):
    """Export a gated cited deliverable as a .pdf (G6.1b).

    F2b: cap-gated on `release_external` (D5) — exporting a deliverable is sending work OUTSIDE
    the firm, the one tightly-held verb. Juniors/paralegals route up the review chain (F2e);
    partners + senior associates may release. Solo-MP / legacy users are allowed (byte-identical).

    The markdown supplied MUST already have passed the output gate. Same contract as
    /export/docx: `include_citations` toggles numbered footnotes vs a clean client copy.
    Uses fpdf2 (pure-Python, already pinned) — no native GTK/Pango system libs required.
    """
    try:
        pdf_bytes = _build_pdf(req)
    except Exception as exc:
        logger.warning("PDF export failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"PDF export failed: {exc}")

    safe_title = "".join(c for c in (req.title or "draft") if c.isalnum() or c in " -_").strip()[:50]
    filename = f"{safe_title or 'draft'}.pdf"
    log_audit(sb, "export.draft", "draft", safe_title,
              {"format": "pdf", "include_citations": req.include_citations})
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/conversations/{conversation_id}/export")
async def export_conversation(
    conversation_id: str,
    format: Literal["md", "pdf"] = "md",
    sb=Depends(get_current_user),
    _cap=Depends(require_cap("release_external")),
):
    """Export a conversation as Markdown or PDF. F2b: cap-gated on `release_external` (D5).

    Query params:
        format: 'md' or 'pdf' (default: 'md')
    """
    # Fetch conversation metadata
    convs = sb.get_conversations()
    conv = next((c for c in convs if c["id"] == conversation_id), None)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    title = conv.get("title", "Untitled Conversation")

    # Fetch messages
    messages = sb.get_messages(conversation_id)
    if not messages:
        raise HTTPException(status_code=404, detail="No messages in this conversation")

    if format == "md":
        md_content = _format_conversation_md(title, messages)
        safe_title = "".join(c for c in title if c.isalnum() or c in " -_").strip()[:50]
        filename = f"{safe_title or 'conversation'}.md"
        log_audit(sb, "export.conversation", "conversation", conversation_id, {"format": "md", "title": title})
        return StreamingResponse(
            io.BytesIO(md_content.encode("utf-8")),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    elif format == "pdf":
        pdf_bytes = _format_conversation_pdf(title, messages)
        safe_title = "".join(c for c in title if c.isalnum() or c in " -_").strip()[:50]
        filename = f"{safe_title or 'conversation'}.pdf"
        log_audit(sb, "export.conversation", "conversation", conversation_id, {"format": "pdf", "title": title})
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")


@router.post("/tables/export")
async def export_table_xlsx(
    req: TableExportRequest,
    sb=Depends(get_current_user),
    _cap=Depends(require_cap("release_external")),
):
    """Export an Analyst table answer as an .xlsx workbook (Phase 4.3, §4b step 4).

    F2b: cap-gated on `release_external` (D5).

    Two sheets: the source table grid and the computed results (value + formula +
    source-cell trace). Computed columns are preserved and labelled, so every
    exported number still traces to a source cell or a shown formula.
    """
    try:
        xlsx_bytes = _build_xlsx(req)
    except Exception as exc:
        logger.warning("XLSX export failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"XLSX export failed: {exc}")

    safe_title = "".join(c for c in (req.title or "table") if c.isalnum() or c in " -_").strip()[:50]
    filename = f"{safe_title or 'table'}.xlsx"
    log_audit(sb, "export.table", "table", safe_title, {"format": "xlsx"})
    return StreamingResponse(
        io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/export/docx")
async def export_draft_docx(
    req: DocxExportRequest,
    sb=Depends(get_current_user),
    _cap=Depends(require_cap("release_external")),
):
    """Export a gated cited deliverable (memo/summary/draft) as a .docx (G6.1).

    F2b: cap-gated on `release_external` (D5).

    The markdown supplied MUST already have passed the output gate — export never
    re-gates and never adds a figure. `include_citations` toggles numbered endnotes
    (the marker `[... p.N]` → superscript ref + Sources section) vs a clean client
    copy with markers stripped. Body numbers are identical either way.
    """
    try:
        docx_bytes = _build_docx(req)
    except Exception as exc:
        logger.warning("DOCX export failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {exc}")

    safe_title = "".join(c for c in (req.title or "draft") if c.isalnum() or c in " -_").strip()[:50]
    filename = f"{safe_title or 'draft'}.docx"
    log_audit(sb, "export.draft", "draft", safe_title,
              {"format": "docx", "include_citations": req.include_citations})
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/export/redline-docx")
async def export_redline_docx(
    req: RedlineExportRequest,
    sb=Depends(get_current_user),
    _cap=Depends(require_cap("release_external")),
):
    """Export a redline result as a tracked-changes .docx (G6.3).

    F2b: cap-gated on `release_external` (D5).

    Each deviation finding renders as a clause pair: the original (struck-through, red)
    and the suggested edit (underlined, green) — the lawyer's redline. Conforming,
    missing, and abstained clauses are summarised at the end. Export is dumb plumbing:
    the agent already grounded the findings; no re-gating, no new intelligence here.
    """
    try:
        docx_bytes = _build_redline_docx(req)
    except Exception as exc:
        logger.warning("Redline DOCX export failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Redline DOCX export failed: {exc}")

    safe_title = "".join(c for c in (req.title or "redline") if c.isalnum() or c in " -_").strip()[:50]
    filename = f"{safe_title or 'redline'}.docx"
    log_audit(sb, "export.redline", "redline", safe_title, {"format": "docx"})
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
